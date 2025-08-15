import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import random
import io
import textwrap
import re

# Optional deps (handled gracefully if missing)
try:
    import google.generativeai as genai  # pip install google-generativeai
except Exception:
    genai = None

try:
    from docx import Document  # pip install python-docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4  # pip install reportlab
    from reportlab.pdfgen import canvas as rl_canvas
except Exception:
    rl_canvas = None

# -----------------------------
# Page config & helpers
# -----------------------------
st.set_page_config(page_title="Auto Research Writer (LLM Edition)", page_icon="🧠", layout="wide")
st.markdown("""
<style>
.block-container { padding-top: 1.2rem; padding-bottom: 2rem; }
footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

# -----------------------------
# App settings
# -----------------------------
ARTICLE_TYPES = {
    "Systematic Review": [
        "Abstract",
        "1. Introduction",
        "2. Methods (Search & Selection)",
        "3. Results (Bibliometrics/Findings)",
        "4. Thematic Synthesis",
        "5. Gaps & Future Directions",
        "6. Conclusion",
    ],
    "Empirical Study": [
        "Abstract",
        "1. Introduction",
        "2. Literature Review",
        "3. Data & Methods",
        "4. Results",
        "5. Robustness & Limitations",
        "6. Conclusion",
    ],
    "Policy Brief": [
        "Executive Summary",
        "1. Background",
        "2. Problem Statement",
        "3. Policy Options",
        "4. Recommendations",
        "5. Implementation & Monitoring",
        "6. Conclusion",
    ],
}

DEFAULT_PROMPT_INSTRUCTIONS = (
    "You are an expert academic writer. Write clear, cohesive, and citation-aware prose suitable for a peer-reviewed venue. "
    "Avoid hallucinated statistics. If asserting facts, hedge appropriately or request verification. Use concise paragraphs."
)

# Fallback offline generators
SENTENCE_TEMPLATES = [
    "{topic} has emerged as a central theme across academia, industry, and policy in the last decade.",
    "Despite rapid progress, open questions remain regarding scalability, equity, and long-term sustainability of {topic}.",
    "We synthesize the state of knowledge on {topic}, identify gaps, and outline actionable directions for future research.",
    "Our analysis balances conceptual clarity with empirical rigor, offering a coherent narrative on {topic}.",
    "Findings suggest that targeted investment, rigorous evaluation, and transparent governance are critical enablers for {topic}.",
    "The implications of {topic} extend to environmental, economic, and social dimensions, requiring cross-disciplinary collaboration.",
]
METHOD_TEMPLATES = [
    "We adopt a transparent protocol, preregistering our plan and adhering to established reporting standards.",
    "Inclusion criteria emphasized relevance to {topic}, methodological rigor, and replicability.",
    "Quantitative synthesis was complemented by qualitative thematic analysis to capture nuance in {topic}.",
    "We performed sensitivity checks and triangulated across data sources to mitigate bias.",
]
RESULT_TEMPLATES = [
    "Across scenarios, we observe consistent improvements associated with targeted interventions in {topic}.",
    "Effect sizes indicate practical significance, though heterogeneity suggests contextual dependencies.",
    "Exploratory analyses reveal non-linearities and potential threshold effects relevant to {topic}.",
]
DISCUSS_TEMPLATES = [
    "The results align with prior work but extend it by formalizing assumptions and testing out-of-sample.",
    "We caution against overgeneralization; external validity depends on data quality and institutional capacity.",
    "Future studies should prioritize open data, preregistered designs, and standardized metrics for {topic}.",
]

def _rand_sent(templates, topic, n=4):
    picks = random.sample(templates, k=min(n, len(templates)))
    return " ".join(t.format(topic=topic) for t in picks)

def offline_section(topic: str, section: str) -> str:
    if re.search(r"Abstract|Executive Summary", section, re.I):
        bank = SENTENCE_TEMPLATES
    elif re.search(r"Method|Selection|Data", section, re.I):
        bank = METHOD_TEMPLATES
    elif re.search(r"Result", section, re.I):
        bank = RESULT_TEMPLATES
    elif re.search(r"Discussion|Gaps|Future", section, re.I):
        bank = DISCUSS_TEMPLATES
    else:
        bank = SENTENCE_TEMPLATES + DISCUSS_TEMPLATES
    text = _rand_sent(bank, topic, n=4)
    for _ in range(3):
        text += " " + random.choice(bank).format(topic=topic)
    return text

# -----------------------------
# LLM (Gemini) helpers
# -----------------------------
def make_gemini_model(api_key: str, model_name: str = "gemini-1.5-flash"):
    if genai is None:
        raise RuntimeError("google-generativeai not installed. Add it to requirements.txt")
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(model_name)

def compose_prompt(topic: str, article_type: str, outline: list, words_per_section: int,
                   extra_instructions: str, want_refs: bool):
    outline_block = "\n".join(f"## {i+1}. {sec}" for i, sec in enumerate(outline))
    refs_guidance = (
        "Include a final section 'References' with plausible APA-style citations (author, year, title, venue). "
        "Do NOT invent DOIs or URLs; if uncertain, mark as 'retrieval needed'."
    ) if want_refs else ""
    return f"""
{DEFAULT_PROMPT_INSTRUCTIONS}
{extra_instructions}

Write a {article_type} on the topic: "{topic}".
Target ~{words_per_section} words per section. Keep paragraphs short and readable.
Use the following outline exactly and write content under each header:
{outline_block}

Tone: scholarly but accessible. Avoid filler. Where evidence is uncertain, say so explicitly.
{refs_guidance}

Return markdown with level-2 headers matching the outline, plus a top-level H1 title and an author line placeholder.
""".strip()

def generate_with_gemini(api_key: str, topic: str, article_type: str, outline: list,
                         words_per_section: int, extra_instructions: str,
                         want_refs: bool, model_name: str):
    prompt = compose_prompt(topic, article_type, outline, words_per_section, extra_instructions, want_refs)
    model = make_gemini_model(api_key, model_name)
    resp = model.generate_content(prompt)
    return resp.text or ""

# -----------------------------
# Export helpers (DOCX + PDF)
# -----------------------------
def parse_markdown_sections(md_text: str):
    title = "Untitled"
    authors = "Anonymous"
    lines = md_text.splitlines()
    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
    for ln in lines[:10]:
        if ln.lower().startswith("**authors**") or ln.lower().startswith("authors:"):
            authors = re.sub(r"(?i)\*\*?authors\*\*?:?\s*", "", ln).strip()
    sections = []
    cur = None
    buf = []
    for ln in lines:
        if ln.startswith("## "):
            if cur is not None:
                sections.append((cur, "\n".join(buf).strip()))
            cur = ln[3:].strip()
            buf = []
        else:
            buf.append(ln)
    if cur is not None:
        sections.append((cur, "\n".join(buf).strip()))
    return title, authors, sections

def export_docx(md_text: str) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx not installed")
    title, authors, sections = parse_markdown_sections(md_text)
    doc = Document()
    doc.add_heading(title, 0)
    p = doc.add_paragraph(f"Authors: {authors}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d"))
    doc.add_paragraph("")
    for sec, body in sections:
        doc.add_heading(sec, level=1)
        for para in body.split("\n\n"):
            doc.add_paragraph(para)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_pdf_basic(md_text: str) -> bytes:
    if rl_canvas is None:
        raise RuntimeError("reportlab not installed")
    title, authors, sections = parse_markdown_sections(md_text)
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    def draw_wrapped(text, x, y, max_width):
        wrapper = textwrap.TextWrapper(width=95)
        for line in text.split("\n"):
            for sub in wrapper.wrap(line):
                nonlocal_y[0] -= 14
                if nonlocal_y[0] < 72:
                    c.showPage(); nonlocal_y[0] = height - 72
                c.drawString(x, nonlocal_y[0], sub)

    nonlocal_y = [height - 72]
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(width/2, nonlocal_y[0], title)
    c.setFont("Helvetica", 10)
    nonlocal_y[0] -= 24
    c.drawCentredString(width/2, nonlocal_y[0], f"Authors: {authors}")
    nonlocal_y[0] -= 24
    c.drawCentredString(width/2, nonlocal_y[0], datetime.now().strftime("%Y-%m-%d"))

    c.setFont("Helvetica", 11)
    for sec, body in sections:
        nonlocal_y[0] -= 28
        if nonlocal_y[0] < 100:
            c.showPage(); nonlocal_y[0] = height - 72; c.setFont("Helvetica", 11)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, nonlocal_y[0], sec)
        c.setFont("Helvetica", 11)
        nonlocal_y[0] -= 10
        draw_wrapped(body, 72, nonlocal_y[0], width - 144)

    c.showPage()
    c.save()
    return buf.getvalue()

# -----------------------------
# UI
# -----------------------------
st.title("🧠 Auto Research Writer — Gemini + Exports")
st.caption("Generate academically-styled articles using Google Gemini or an offline template. Export to DOCX/PDF.")

left, right = st.columns([1.5, 1])
with left:
    topic = st.text_input("Topic / Chủ đề", value="Green growth and sustainability")
    article_type = st.selectbox("Article type", list(ARTICLE_TYPES.keys()), index=0)
    outline = ARTICLE_TYPES[article_type]
    words_per_section = st.slider("~Words per section", 120, 600, 220, step=20)
    want_refs = st.checkbox("Ask model to include 'References' section", value=True)
    extra_instructions = st.text_area("Extra instructions (optional)", value="Use neutral, precise language and avoid unsupported claims.")

with right:
    mode = st.radio("Generation mode", ["Gemini (LLM)", "Offline Template"], index=0)
    model_name = st.selectbox("Gemini model", ["gemini-1.5-flash", "gemini-1.5-pro"], index=0)
    api_key = st.text_input("GEMINI_API_KEY (or use st.secrets)", type="password", value=st.secrets.get("GEMINI_API_KEY", ""))
    seed = st.number_input("Random seed (optional)", value=0, step=1)
    include_demo_chart = st.checkbox("Add demo chart (synthetic)", value=False)

if seed:
    random.seed(int(seed)); np.random.seed(int(seed))

st.divider()

if st.button("🚀 Generate Article", type="primary"):
    if mode.startswith("Gemini"):
        if not api_key:
            st.error("Please provide GEMINI_API_KEY (Settings ➜ Secrets on Streamlit Cloud). Falling back to offline template.")
            mode = "Offline Template"

    if mode.startswith("Gemini"):
        try:
            md = generate_with_gemini(api_key, topic, article_type, outline, words_per_section,
                                      extra_instructions, want_refs, model_name)
            if not md.strip():
                raise RuntimeError("Empty response from Gemini")
        except Exception as e:
            st.warning(f"Gemini error: {e}. Using offline generator instead.")
            md = None
    else:
        md = None

    if md is None:
        title = f"A systematic literature review of {topic.lower()}" if article_type == "Systematic Review" else f"On {topic}"
        parts = [f"# {title.title()}", f"**Authors:** Placeholder", f"**Date:** {datetime.now().strftime('%Y-%m-%d')}", ""]
        for sec in outline:
            parts.append(f"## {sec}")
            parts.append(offline_section(topic, sec))
            parts.append("")
        if want_refs:
            parts.append("## References\n(Will be curated; retrieval needed)")
        md = "\n".join(parts)

    st.markdown(md)

    if include_demo_chart:
        st.subheader("Figure: Demonstration Time Series")
        x = pd.date_range(datetime.now().date().replace(day=1), periods=36, freq="MS")
        y1 = np.cumsum(np.random.randn(len(x))) + 10
        y2 = y1 + np.random.randn(len(x)) * 0.5 + 2
        fig_df = pd.DataFrame({"Date": x, "Baseline": y1, "Proposed": y2}).set_index("Date")
        st.line_chart(fig_df, use_container_width=True)
        st.caption("Synthetic data for illustrative purposes only.")

    st.divider()
    st.download_button("⬇️ Download Markdown", data=md, file_name="article.md", mime="text/markdown")

    try:
        docx_bytes = export_docx(md)
        st.download_button("⬇️ Download DOCX", data=docx_bytes,
            file_name="article.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.info(f"DOCX export unavailable: {e}")

    try:
        pdf_bytes = export_pdf_basic(md)
        st.download_button("⬇️ Download PDF", data=pdf_bytes, file_name="article.pdf", mime="application/pdf")
    except Exception as e:
        st.info(f"PDF export unavailable: {e}")

else:
    st.info("Nhập chủ đề, chọn chế độ, rồi bấm **Generate Article**.")

with st.expander("ℹ️ Setup (GitHub + Streamlit Cloud)"):
    st.markdown(
        """

